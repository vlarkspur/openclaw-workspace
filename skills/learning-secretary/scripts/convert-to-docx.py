#!/usr/bin/env python3
"""
Markdown 转 Word 脚本
用于将学习周报从 Markdown 格式转换为 .docx 格式

用法：
    python3 convert-to-docx.py input.md [output.docx]

如果未指定输出文件，自动生成同名的 .docx 文件
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 缺少 python-docx 库，请先安装：")
    print("   pip3 install python-docx")
    sys.exit(1)


def markdown_to_docx(md_path, docx_path=None):
    """简单的 Markdown 转 Word 转换器"""
    
    if docx_path is None:
        docx_path = Path(md_path).with_suffix('.docx')
    
    doc = Document()
    
    # 读取 Markdown 文件
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_list = []
    in_code_block = False
    code_content = []
    
    for line in lines:
        line_stripped = line.rstrip('\n')
        
        # 处理代码块
        if line_stripped.startswith('```'):
            if in_code_block:
                # 结束代码块
                code_para = doc.add_paragraph('\n'.join(code_content))
                code_para.style = 'No Spacing'
                code_para.runs[0].font.name = 'Courier New'
                code_para.runs[0].font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        
        if in_code_block:
            code_content.append(line_stripped)
            continue
        
        # 处理标题
        if line_stripped.startswith('# '):
            doc.add_heading(line_stripped[2:], level=1)
        elif line_stripped.startswith('## '):
            doc.add_heading(line_stripped[3:], level=2)
        elif line_stripped.startswith('### '):
            doc.add_heading(line_stripped[4:], level=3)
        # 处理列表
        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
            item = line_stripped[2:]
            if not current_list:
                current_list = doc.add_paragraph(style='List Bullet')
                current_list.runs[0].text = item
            else:
                para = doc.add_paragraph(style='List Bullet')
                para.runs[0].text = item
        # 处理表格行（简单处理）
        elif line_stripped.startswith('|'):
            # 跳过表格，因为简单转换不支持复杂表格
            pass
        # 处理普通段落
        elif line_stripped.strip():
            doc.add_paragraph(line_stripped)
            current_list = None
        else:
            current_list = None
    
    # 保存文档
    doc.save(docx_path)
    print(f"✅ 已生成：{docx_path}")
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法：python3 convert-to-docx.py input.md [output.docx]")
        sys.exit(1)
    
    md_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(md_path):
        print(f"❌ 文件不存在：{md_path}")
        sys.exit(1)
    
    markdown_to_docx(md_path, docx_path)
